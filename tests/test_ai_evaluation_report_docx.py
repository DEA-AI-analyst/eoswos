from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from ai_evaluation_report import build_canonical_report_context
from ai_evaluation_report_docx import (
    _price_consistency_sentence,
    _source_note_display,
    build_report_docx,
    report_filename,
)
from test_ai_evaluation_report import _result, _submitted


KNOWN_SOURCE_NOTE = (
    "DART CFS assets, liabilities, and non-controlling interests with KRX listed shares."
)
KOREAN_SOURCE_NOTE = (
    "DART 연결재무제표의 자산·부채·비지배지분과 KRX 상장주식수를 기준으로 산정"
)


def _document_text(document: Document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        chunks.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(chunks)


def _body_runs(document: Document):
    for paragraph in document.paragraphs:
        yield from paragraph.runs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield from paragraph.runs


def _all_cells_centered(table) -> bool:
    return all(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )


def test_docx_is_deterministic_a4_report_with_confirmed_values() -> None:
    result = _result()
    result["bps_provenance"]["source_note"] = KNOWN_SOURCE_NOTE
    context = build_canonical_report_context(
        result,
        submitted_input=_submitted(),
        model_mode="FROZEN_REFERENCE",
    )
    output = build_report_docx(
        context,
        ai_commentary=(
            "확정 결과의 상대적 검토 우선순위를 확인했습니다.\n"
            "p_M과 s_M의 질적 해석을 함께 확인해야 합니다.\n"
            "추가 재무 및 시장정보를 검토해 주세요."
        ),
        reviewer_comment="담당자 최종 검토의견입니다.",
    )
    assert output.startswith(b"PK")
    document = Document(BytesIO(output))
    section = document.sections[0]
    assert abs(section.page_width.mm - 210) < 0.1
    assert abs(section.page_height.mm - 297) < 0.1
    text = _document_text(document)
    for expected in (
        "EosWos AI 평가보고서",
        "AI Evaluation Report",
        "신규평가",
        "M4",
        "34",
        "1,261",
        "0.682",
        "DART 재무 + 검증된 KRX 상장주",
        "2026-06-30 / CFS",
        "FROZEN_REFERENCE",
        KOREAN_SOURCE_NOTE,
        "담당자 최종 검토의견입니다.",
        "E2 | Score 0.880 | Rank 400",
        "1D·1W·1M 가격기준에서 M Grade가 M4로 동일하게 유지됩니다.",
    ):
        assert expected in text

    key_results = next(
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["M Grade", "M Score", "M Rank", "Final Score"]
    )
    assert key_results.cell(1, 1).text == "34"
    assert key_results.cell(1, 3).text == "0.682"
    assert abs(key_results.cell(1, 0).paragraphs[0].runs[0].font.size.pt - 15.0) < 0.01
    for index in (1, 2, 3):
        assert abs(key_results.cell(1, index).paragraphs[0].runs[0].font.size.pt - 10.0) < 0.01
    for index in range(4):
        margins = key_results.cell(1, index)._tc.get_or_add_tcPr().find(qn("w:tcMar"))
        assert margins is not None
        assert margins.find(qn("w:top")).get(qn("w:w")) == "45"
        assert margins.find(qn("w:bottom")).get(qn("w:w")) == "45"

    provenance = next(
        table for table in document.tables if table.cell(0, 0).text == "평가 데이터"
    )
    grid_widths = [int(column.get(qn("w:w"))) for column in provenance._tbl.tblGrid]
    assert grid_widths == [1500, 4400, 1400, 3580]
    assert sum(grid_widths) <= 10885
    assert provenance.cell(0, 1).text == "DART 재무 + 검증된 KRX 상장주"
    assert provenance.cell(0, 3).text == "2026-06-30 / CFS"
    for row in provenance.rows[:3]:
        for cell in row.cells:
            margins = cell._tc.get_or_add_tcPr().find(qn("w:tcMar"))
            assert margins is not None
            assert margins.find(qn("w:start")).get(qn("w:w")) == "55"
            assert margins.find(qn("w:end")).get(qn("w:w")) == "55"

    axes = next(
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells] == ["e_M", "p_M", "s_M"]
    )
    assert _all_cells_centered(axes)
    assert [cell.text for cell in axes.rows[1].cells] == [
        "구조적 효율성",
        "ITM 도달가능성 + PK 강도",
        "First_ITM Timing / 도달속도",
    ]

    price = next(
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["가격기준", "M Grade", "e_M", "p_M", "s_M"]
    )
    assert _all_cells_centered(price)
    legend = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "축 열 표기: Grade | Score | Rank"
    )
    assert legend.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    price_grid_widths = [int(column.get(qn("w:w"))) for column in price._tbl.tblGrid]
    assert price_grid_widths == [1157, 1157, 2922, 2922, 2922]
    assert sum(price_grid_widths) == 11080
    assert abs(price_grid_widths[0] * 2.54 / 1440 - 2.04) < 0.001
    assert abs(price_grid_widths[1] * 2.54 / 1440 - 2.04) < 0.001
    assert len(set(price_grid_widths[2:])) == 1
    table_width = price._tbl.tblPr.find(qn("w:tblW"))
    assert table_width is not None
    assert table_width.get(qn("w:type")) == "dxa"
    assert int(table_width.get(qn("w:w"))) == 11080
    layout = price._tbl.tblPr.find(qn("w:tblLayout"))
    assert layout is not None
    assert layout.get(qn("w:type")) == "fixed"
    assert price.autofit is False
    for row in price.rows:
        row_widths = [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells]
        assert row_widths == price_grid_widths

    review = document.tables[-1]
    assert len(review.rows) == 1 and len(review.columns) == 1
    assert review.cell(0, 0).text == "담당자 최종 검토의견입니다."
    borders = review.cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
    assert borders is not None

    assert "EOSWOS" not in text
    assert "Scoring Source" not in text
    assert "Financial Entity" not in text
    assert "Source Note" not in text
    assert KNOWN_SOURCE_NOTE not in text
    for run in _body_runs(document):
        if not run.text:
            continue
        assert round(run.font.size.pt, 1) in {10.0, 15.0, 16.5}
    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    assert abs(run.font.size.pt - 8.0) < 0.01
    assert "본 보고서는 AI 기반 의사결정 지원자료이며 최종 투자판단은 별도의 검토를 통해 이루어집니다." in text
    assert "M Grade는 상대적 검토 우선순위를 나타내며" not in text
    assert "999999" not in text


def test_price_summary_is_deterministic_for_equal_and_mixed_grades() -> None:
    assert _price_consistency_sentence((("1D", "M2"), ("1W", "M2"), ("1M", "M2"))) == (
        "1D·1W·1M 가격기준에서 M Grade가 M2로 동일하게 유지됩니다."
    )
    assert _price_consistency_sentence((("1D", "M2"), ("1W", "M3"), ("1M", "M2"))) == (
        "가격기준별 M Grade가 1D M2 · 1W M3 · 1M M2로 달라 기준별 결과 차이를 함께 검토해야 합니다."
    )


def test_source_note_translation_is_exact_and_unknown_text_is_preserved() -> None:
    assert _source_note_display(KNOWN_SOURCE_NOTE) == KOREAN_SOURCE_NOTE
    unknown = "Approved future source note"
    assert _source_note_display(unknown) == unknown


def test_max_length_layout_preserves_full_text_at_ten_points() -> None:
    context = build_canonical_report_context(
        _result(),
        submitted_input=_submitted(),
        model_mode="FROZEN_REFERENCE",
    )
    ai_text = ("확정 사실 범위에서만 해석합니다. " * 100)[:1400]
    reviewer_text = ("담당자 검토의견을 기록합니다. " * 100)[:790]

    output = build_report_docx(
        context,
        ai_commentary=ai_text,
        reviewer_comment=reviewer_text,
    )
    document = Document(BytesIO(output))

    assert ai_text in [paragraph.text for paragraph in document.paragraphs]
    assert document.tables[-1].cell(0, 0).text == reviewer_text
    ai_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == ai_text)
    assert all(abs(run.font.size.pt - 10.0) < 0.01 for run in ai_paragraph.runs if run.text)
    review_runs = document.tables[-1].cell(0, 0).paragraphs[0].runs
    assert all(abs(run.font.size.pt - 10.0) < 0.01 for run in review_runs if run.text)
    assert abs(document.sections[0].top_margin.mm - 5) < 0.1
    assert abs(document.sections[0].bottom_margin.mm - 6) < 0.1

def test_report_filename_is_safe_and_docx_only() -> None:
    context = build_canonical_report_context(_result(), submitted_input=_submitted())
    filename = report_filename(context)
    assert filename.startswith("EosWos_AI_평가보고서_현대건설_")
    assert filename.endswith(".docx")
    assert "/" not in filename and "\\" not in filename


def test_monitoring_uses_same_template_with_optional_fields() -> None:
    result = _result()
    result.update(
        {
            "Monitoring_Date": "2026-06-30",
            "Actual_Issue_Date": "2024-07-05",
            "FS_Target_Date": "2026-06-30",
            "Original_TTM_years": 5.0,
            "Rebased_TTM_years_raw": 3.01,
        }
    )
    context = build_canonical_report_context(result, model_mode="FROZEN_REFERENCE")
    output = build_report_docx(context, ai_commentary="사후관리 결과를 확인했습니다.", reviewer_comment="")
    document = Document(BytesIO(output))
    text = _document_text(document)
    assert "사후관리 재평가" in text
    assert "Actual 2024-07-05" in text
    assert "Rebased 3.01" in text
    assert "별도의 독립 Monitoring Model을 의미하지 않습니다" in text
    assert document.tables[-1].cell(0, 0).text == "미작성"