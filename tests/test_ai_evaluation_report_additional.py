import copy
from io import BytesIO

from docx import Document

from ai_evaluation_report import build_canonical_report_context
from ai_evaluation_report_docx import (
    ReportDocumentError,
    build_report_docx,
)
from test_ai_evaluation_report import _result, _submitted
from test_ai_evaluation_report_streamlit import _confirmed_result, _show_confirmed
from test_streamlit_chat_routing import _app


def _document_text(document: Document) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_missing_provenance_is_not_invented() -> None:
    result = _result()
    result.pop("bps_provenance")

    context = build_canonical_report_context(
        result,
        submitted_input=_submitted(),
        model_mode="FROZEN_REFERENCE",
    )

    provenance = context["provenance"]
    assert provenance == {"model_mode": "FROZEN_REFERENCE"}
    assert "DART" not in str(provenance)
    assert "KRX" not in str(provenance)


def test_docx_failure_preserves_confirmed_evaluation(monkeypatch) -> None:
    at, _ = _app(monkeypatch)
    confirmed = _confirmed_result()
    before = copy.deepcopy(confirmed)
    _show_confirmed(at, confirmed)
    next(button for button in at.button if button.label == "AI 평가보고서").click().run(timeout=10)

    def fail(*_args, **_kwargs):
        raise ReportDocumentError("private traceback C:/secret/report.py")

    monkeypatch.setattr("ai_evaluation_report_ui.build_report_docx", fail)
    at.run(timeout=10)

    assert not at.exception
    assert at.session_state["current_evaluation"] == before
    public = " ".join(str(item.value) for item in at.error)
    assert "AI 평가보고서 DOCX를 생성하지 못했습니다" in public
    assert "private" not in public
    assert "C:/" not in public


def test_price_consistency_uses_only_provided_bases() -> None:
    result = _result()
    result["price_basis"].pop("1M")
    result["price_basis"]["1W"]["m_grade"] = "M3"
    context = build_canonical_report_context(result, submitted_input=_submitted())

    output = build_report_docx(
        context,
        ai_commentary="제공된 가격기준 결과를 확인했습니다.",
        reviewer_comment="",
    )
    text = _document_text(Document(BytesIO(output)))

    assert "가격기준별 M Grade가 1D M4 · 1W M3" in text
    assert "1M M" not in text
