"""Deterministic one-page DOCX renderer for the EosWos AI evaluation report."""

from __future__ import annotations

import io
import re
from typing import Any, Mapping, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


NAVY = "17365D"
BLUE = "2F75B5"
PALE_BLUE = "EAF1F8"
PALE_GRAY = "F4F6F8"
MID_GRAY = "6B7280"
WHITE = "FFFFFF"
BLACK = "111827"
FONT = "Malgun Gothic"
BODY_FONT_SIZE = 10.0
FOOTER_FONT_SIZE = 8.0
REPORT_DISCLAIMER = (
    "M Grade는 상대적 검토 우선순위를 나타내며 개별 성공확률이나 투자승인·부결을 의미하지 "
    "않습니다. 본 보고서는 AI 기반 의사결정 지원자료이며 최종 투자판단은 별도의 검토를 통해 이루어집니다."
)
MONITORING_DISCLAIMER = (
    "Monitoring M Grade는 Frozen Model을 Rebased 방식으로 적용한 사후관리 평가결과이며, "
    "별도의 독립 Monitoring Model을 의미하지 않습니다."
)


class ReportDocumentError(RuntimeError):
    """Public-safe DOCX rendering failure."""


def build_report_docx(
    canonical_context: Mapping[str, Any],
    *,
    ai_commentary: str,
    reviewer_comment: str,
) -> bytes:
    """Render report data without recalculating or asking an LLM for any value."""

    try:
        document = Document()
        _configure_document(document)
        _add_header(document, canonical_context)
        _add_overview(document, canonical_context)
        _add_key_results(document, canonical_context)
        _add_axes(document, canonical_context)
        _add_price_basis(document, canonical_context)
        _add_provenance(document, canonical_context)
        _add_ai_commentary(document, ai_commentary or "미작성")
        _add_review_commentary(document, reviewer_comment or "미작성")
        _add_footer(document, canonical_context)
        if _requires_compact_layout(ai_commentary, reviewer_comment):
            _apply_compact_layout(document)
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception as exc:
        raise ReportDocumentError("AI 평가보고서 DOCX를 생성하지 못했습니다.") from exc


def report_filename(canonical_context: Mapping[str, Any]) -> str:
    common = _mapping(canonical_context.get("common_info"))
    company = str(common.get("company") or "evaluation").strip()
    request_id = str(canonical_context.get("request_id") or "report").strip()
    safe_company = re.sub(r"[^0-9A-Za-z가-힣_-]+", "_", company).strip("_") or "evaluation"
    safe_request = re.sub(r"[^0-9A-Za-z_-]+", "", request_id)[:16] or "report"
    return f"EosWos_AI_평가보고서_{safe_company}_{safe_request}.docx"


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(13)
    section.left_margin = Mm(9)
    section.right_margin = Mm(9)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(6)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal.font.size = Pt(BODY_FONT_SIZE)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0


def _requires_compact_layout(ai_commentary: str, reviewer_comment: str) -> bool:
    return len(str(ai_commentary or "")) + len(str(reviewer_comment or "")) > 900


def _apply_compact_layout(document: Document) -> None:
    """Preserve 10 pt text while tightening only unusually long one-page drafts."""

    section = document.sections[0]
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(6)
    section.header_distance = Mm(3)
    section.footer_distance = Mm(2.5)

    in_ai_commentary = False
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value == "6. AI 평가의견":
            in_ai_commentary = True
        elif value == "7. 최종 검토의견":
            in_ai_commentary = False
        if re.match(r"^[1-7]\. ", value):
            paragraph.paragraph_format.space_before = Pt(1.5)
            paragraph.paragraph_format.space_after = Pt(0.5)
            paragraph.paragraph_format.line_spacing = Pt(11.2)
        elif value == "AI Evaluation Report   |   평가유형: 신규평가" or value.startswith(
            "AI Evaluation Report   |   평가유형:"
        ):
            paragraph.paragraph_format.space_after = Pt(0.5)
        elif in_ai_commentary:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0.35)
            paragraph.paragraph_format.line_spacing = Pt(10.5)
        else:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _set_cell_margins(cell, top=20, bottom=20, start=55, end=55)
                for paragraph in cell.paragraphs:
                    largest = max(
                        (run.font.size.pt for run in paragraph.runs if run.font.size),
                        default=BODY_FONT_SIZE,
                    )
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = Pt(16.5 if largest >= 15 else 10.8)

    for footer in (section.footer,):
        for paragraph in footer.paragraphs:
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(8.5)

def _add_header(document: Document, context: Mapping[str, Any]) -> None:
    evaluation_type = str(context.get("evaluation_type") or "평가")
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(title, "EosWos AI 평가보고서", size=16.5, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(2)
    _run(subtitle, "AI Evaluation Report", size=BODY_FONT_SIZE, bold=True, color=BLUE)
    _run(subtitle, f"   |   평가유형: {evaluation_type}", size=BODY_FONT_SIZE, color=MID_GRAY)
    _paragraph_rule(subtitle, color=BLUE, size="10")


def _add_overview(document: Document, context: Mapping[str, Any]) -> None:
    _section_heading(document, "1. 평가 개요")
    common = _mapping(context.get("common_info"))
    monitoring = _mapping(context.get("monitoring_fields"))
    rows = [
        ("평가대상", common.get("company"), "상품유형", common.get("product_type")),
        ("기초자산", common.get("underlying_company") or common.get("stock_code"), "평가기준일", common.get("price_basis_date")),
        ("발행조건", _issue_condition_text(common), "신용등급", common.get("credit_rating")),
    ]
    if not monitoring:
        rows.append(
            (
                "Issue Date",
                common.get("issue_date"),
                "TTM",
                _joined_values(
                    ("계약", _number(common.get("contract_ttm_years"), 2)),
                    ("모델", _number(common.get("model_ttm_years"), 2)),
                ),
            )
        )
    if monitoring:
        rows.append(
            (
                "기간정보",
                _joined_values(
                    ("Actual", monitoring.get("actual_issue_date")),
                    ("Monitoring", monitoring.get("monitoring_date")),
                    ("FS", monitoring.get("fs_target_date")),
                ),
                "TTM",
                _joined_values(
                    ("Original", monitoring.get("original_ttm_years")),
                    ("Rebased", monitoring.get("rebased_ttm_years")),
                    ("Model", monitoring.get("model_ttm_years")),
                ),
            )
        )
    _label_value_table(document, rows, widths=(1900, 3500, 1700, 3780))


def _add_key_results(document: Document, context: Mapping[str, Any]) -> None:
    _section_heading(document, "2. 핵심 평가결과")
    result = _mapping(context.get("evaluation_result"))
    widths = (3200, 2626, 2626, 2628)
    table = document.add_table(rows=2, cols=4)
    _configure_table(table, widths)
    headers = ("M Grade", "M Score", "M Rank", "Final Score")
    values = (
        result.get("m_grade"),
        _number(result.get("m_score"), 0),
        _rank(result.get("m_rank")),
        _number(result.get("final_score"), 3),
    )
    for index, label in enumerate(headers):
        _cell_text(
            table.cell(0, index),
            label,
            bold=True,
            color=WHITE,
            size=BODY_FONT_SIZE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _shade(table.cell(0, index), NAVY)
        value_size = 15.0 if index == 0 else BODY_FONT_SIZE
        _cell_text(
            table.cell(1, index),
            values[index],
            bold=True,
            color=NAVY,
            size=value_size,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_margins(table.cell(1, index), top=80, bottom=80, start=85, end=85)
        _shade(table.cell(1, index), PALE_BLUE if index == 0 else WHITE)

def _add_axes(document: Document, context: Mapping[str, Any]) -> None:
    _section_heading(document, "3. 핵심 평가축")
    axes = _mapping(context.get("axis_results"))
    table = document.add_table(rows=3, cols=3)
    _configure_table(table, (3693, 3693, 3694))
    definitions = (
        ("e_M", "구조적 효율성", axes.get("e_M")),
        ("p_M", "ITM 도달가능성 + PK 강도", axes.get("p_M")),
        ("s_M", "First_ITM Timing / 도달속도", axes.get("s_M")),
    )
    for index, (name, description, raw_axis) in enumerate(definitions):
        axis = _mapping(raw_axis)
        _cell_text(
            table.cell(0, index),
            name,
            bold=True,
            color=WHITE,
            size=BODY_FONT_SIZE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _shade(table.cell(0, index), BLUE)
        _cell_text(
            table.cell(1, index),
            description,
            bold=True,
            color=NAVY,
            size=BODY_FONT_SIZE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _shade(table.cell(1, index), PALE_BLUE)
        value = _joined_values(
            ("", axis.get("grade")),
            ("Score", _number(axis.get("score"), 3)),
            ("Rank", _rank(axis.get("rank"))),
            separator=" | ",
        )
        _cell_text(
            table.cell(2, index),
            value,
            bold=True,
            color=BLACK,
            size=BODY_FONT_SIZE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

def _add_price_basis(document: Document, context: Mapping[str, Any]) -> None:
    _section_heading(document, "4. 가격기준별 평가")
    rows = context.get("price_basis_results")
    rows = rows if isinstance(rows, Sequence) and not isinstance(rows, (str, bytes)) else []
    legend = document.add_paragraph()
    legend.paragraph_format.space_before = Pt(0)
    legend.paragraph_format.space_after = Pt(0.8)
    _run(legend, "축 열 표기: Grade | Score | Rank", size=BODY_FONT_SIZE, color=MID_GRAY)
    widths = (1200, 1700, 2726, 2726, 2728)
    table = document.add_table(rows=1, cols=5)
    _configure_table(table, widths)
    for index, label in enumerate(("가격기준", "M Grade", "e_M", "p_M", "s_M")):
        _cell_text(
            table.cell(0, index),
            label,
            bold=True,
            color=WHITE,
            size=BODY_FONT_SIZE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _shade(table.cell(0, index), NAVY)
    basis_grades: list[tuple[str, str]] = []
    for raw_row in rows:
        row = _mapping(raw_row)
        cells = table.add_row().cells
        _apply_row_widths(cells, widths)
        values = (
            row.get("price_basis"),
            row.get("m_grade"),
            _axis_compact(row.get("e_M")),
            _axis_compact(row.get("p_M")),
            _axis_compact(row.get("s_M")),
        )
        basis = str(row.get("price_basis") or "")
        basis_grades.append((basis, str(row.get("m_grade") or "")))
        for index, value in enumerate(values):
            _cell_text(
                cells[index],
                value,
                bold=index in {0, 1},
                size=BODY_FONT_SIZE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
    interpretation = _price_consistency_sentence(basis_grades)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1.2)
    paragraph.paragraph_format.space_after = Pt(0)
    _run(paragraph, interpretation, size=BODY_FONT_SIZE, color=MID_GRAY)

def _add_provenance(document: Document, context: Mapping[str, Any]) -> None:
    _section_heading(document, "5. 데이터 및 평가근거")
    provenance = _mapping(context.get("provenance"))
    widths = (1850, 3660, 1550, 4020)
    rows = [
        (
            "평가 데이터",
            provenance.get("scoring_source_label") or provenance.get("scoring_source"),
            "재무제표",
            _financial_statement_display(
                provenance.get("fs_selected_date"),
                provenance.get("fs_type"),
            ),
        ),
        (
            "재무정보 기준",
            provenance.get("financial_entity_label") or provenance.get("financial_entity"),
            "BPS 기준",
            provenance.get("status_label") or provenance.get("status"),
        ),
        (
            "모델 기준",
            provenance.get("model_mode"),
            "평가 ID",
            context.get("request_id"),
        ),
    ]
    table = _label_value_table(document, rows, widths=widths)
    source_note = provenance.get("source_note")
    if source_note:
        cells = table.add_row().cells
        _apply_row_widths(cells, widths)
        _cell_text(cells[0], "자료 기준", bold=True, color=NAVY, size=BODY_FONT_SIZE)
        _shade(cells[0], PALE_GRAY)
        merged = cells[1].merge(cells[3])
        _set_cell_width(merged, sum(widths[1:]))
        _cell_text(
            merged,
            _source_note_display(source_note),
            size=BODY_FONT_SIZE,
            color=BLACK,
        )
        _shade(merged, WHITE)

def _add_ai_commentary(document: Document, text: str) -> None:
    _section_heading(document, "6. AI 평가의견")
    lines = _commentary_lines(text)
    for index, line in enumerate(lines):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0 if index == len(lines) - 1 else 0.8)
        paragraph.paragraph_format.line_spacing = 1.05
        _run(paragraph, line, size=BODY_FONT_SIZE, color=BLACK)


def _add_review_commentary(document: Document, text: str) -> None:
    _section_heading(document, "7. 최종 검토의견")
    table = document.add_table(rows=1, cols=1)
    _configure_table(table, (11080,))
    cell = table.cell(0, 0)
    _set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
    _set_cell_border(cell, color="C9D3DF", size="5")
    _shade(cell, PALE_GRAY)
    _cell_text(cell, str(text or "미작성").strip(), size=BODY_FONT_SIZE, color=BLACK)
    cell.paragraphs[0].paragraph_format.line_spacing = 1.05

def _add_footer(document: Document, context: Mapping[str, Any]) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2.5)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    _paragraph_rule(paragraph, color="C9D3DF", size="4", side="top")
    text = REPORT_DISCLAIMER
    if context.get("evaluation_type") == "사후관리 재평가":
        text = f"{text} {MONITORING_DISCLAIMER}"
    _run(paragraph, text, size=FOOTER_FONT_SIZE, color=MID_GRAY)

def _section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3.2)
    paragraph.paragraph_format.space_after = Pt(1.0)
    paragraph.paragraph_format.keep_with_next = True
    _run(paragraph, text, size=BODY_FONT_SIZE, bold=True, color=NAVY)

def _label_value_table(
    document: Document,
    rows: Sequence[tuple[Any, Any, Any, Any]],
    *,
    widths: tuple[int, int, int, int],
):
    table = document.add_table(rows=0, cols=4)
    _configure_table(table, widths)
    for raw_row in rows:
        cells = table.add_row().cells
        _apply_row_widths(cells, widths)
        for index, value in enumerate(raw_row):
            is_label = index in {0, 2}
            _cell_text(
                cells[index],
                value,
                bold=is_label,
                color=NAVY if is_label else BLACK,
                size=BODY_FONT_SIZE,
                blank_if_empty=value == "",
            )
            _shade(cells[index], PALE_GRAY if is_label else WHITE)
    return table

def _configure_table(table, widths: Sequence[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table_element = table._tbl
    properties = table_element.tblPr
    width_element = properties.find(qn("w:tblW"))
    if width_element is None:
        width_element = OxmlElement("w:tblW")
        properties.append(width_element)
    width_element.set(qn("w:type"), "dxa")
    width_element.set(qn("w:w"), str(sum(widths)))
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")

    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        _apply_row_widths(row.cells, widths)


def _apply_row_widths(cells, widths: Sequence[int]) -> None:
    for cell, width in zip(cells, widths):
        _set_cell_width(cell, width)
        _set_cell_margins(cell, top=45, bottom=45, start=75, end=75)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(width))


def _set_cell_margins(
    cell,
    *,
    top: int,
    bottom: int,
    start: int,
    end: int,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        item = margins.find(qn(f"w:{name}"))
        if item is None:
            item = OxmlElement(f"w:{name}")
            margins.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def _set_cell_border(cell, *, color: str, size: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), color)

def _cell_text(
    cell,
    value: Any,
    *,
    bold: bool = False,
    color: str = BLACK,
    size: float = BODY_FONT_SIZE,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    blank_if_empty: bool = False,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 0.92
    text = "" if blank_if_empty and value in (None, "") else _display(value)
    _run(paragraph, text, bold=bold, color=color, size=size)


def _run(paragraph, text: Any, *, size: float, color: str, bold: bool = False):
    run = paragraph.add_run(str(text))
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    return run


def _shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def _paragraph_rule(
    paragraph,
    *,
    color: str,
    size: str,
    side: str = "bottom",
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)
    borders.append(border)

def _issue_condition_text(common: Mapping[str, Any]) -> str:
    return _joined_values(
        ("전환/행사/교환가", _number(common.get("conversion_price"), 0)),
        ("Call", _number(common.get("call_rate"), 2)),
    )


def _price_consistency_sentence(basis_grades: Sequence[tuple[str, str]]) -> str:
    present = [(basis, grade) for basis, grade in basis_grades if basis and grade]
    if len(present) < 2:
        return "가격기준 일관성은 제공된 결과 범위에서 확인해야 합니다."
    bases = "·".join(basis for basis, _ in present)
    grades = [grade for _, grade in present]
    if len(set(grades)) == 1:
        return f"{bases} 가격기준에서 M Grade가 {grades[0]}로 동일하게 유지됩니다."
    details = " · ".join(f"{basis} {grade}" for basis, grade in present)
    return f"가격기준별 M Grade가 {details}로 달라 기준별 결과 차이를 함께 검토해야 합니다."


def _axis_compact(value: Any) -> str:
    axis = _mapping(value)
    return _joined_values(
        ("", axis.get("grade")),
        ("Score", _number(axis.get("score"), 3)),
        ("Rank", _rank(axis.get("rank"))),
        separator=" | ",
    )

def _joined_values(
    *pairs: tuple[str, Any],
    separator: str = " · ",
) -> str:
    parts = []
    for label, value in pairs:
        if value in (None, "", "-"):
            continue
        parts.append(f"{label} {value}".strip())
    return separator.join(parts) if parts else "미제공"


def _financial_statement_display(date: Any, fs_type: Any) -> str:
    values = [str(value).strip() for value in (date, fs_type) if value not in (None, "", "-")]
    return " / ".join(values) if values else "미제공"


def _source_note_display(value: Any) -> str:
    original = str(value or "").strip()
    normalized = re.sub(r"\s+", " ", original).rstrip(".").casefold()
    mapping = {
        "dart cfs assets, liabilities, and non-controlling interests with krx listed shares": (
            "DART 연결재무제표의 자산·부채·비지배지분과 KRX 상장주식수를 기준으로 산정"
        ),
    }
    return mapping.get(normalized, original or "미제공")


def _commentary_lines(value: Any) -> list[str]:
    text = str(value or "미작성").strip()
    lines = [re.sub(r"\s+$", "", line).strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return lines or ["미작성"]

def _number(value: Any, digits: int) -> str:
    if value in (None, ""):
        return "미제공"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if digits == 0:
        return f"{number:,.0f}"
    return f"{number:,.{digits}f}"


def _rank(value: Any) -> str:
    if value in (None, ""):
        return "미제공"
    try:
        return f"{int(float(value)):,}"
    except (TypeError, ValueError):
        return str(value)


def _display(value: Any) -> str:
    return "미제공" if value in (None, "") else str(value)


def _mapping(value: Any) -> Mapping[str, Any]:
    return value if isinstance(value, Mapping) else {}
